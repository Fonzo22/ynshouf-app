from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE   = RGBColor(0x1B, 0x3A, 0x6B)
LIGHT_BLUE  = RGBColor(0x2E, 0x74, 0xB5)

path = r'c:/Users/Fonz/ynshouf-app/ynshouf-app/ynshouf-app/YNSHOUF_Template.docx'
doc = Document(path)

# ── 1 & 2. Paragraph 2: {{report_type}} (runs 0-2) + {{report_sub}} (runs 4-6) ──
p2 = doc.paragraphs[2]
for idx in [0, 1, 2]:
    r = p2.runs[idx]
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DARK_BLUE

for idx in [4, 5, 6]:
    r = p2.runs[idx]
    r.bold = False
    r.font.size = Pt(14)
    r.font.color.rgb = LIGHT_BLUE

# ── 3. Paragraph 4: {{incident_description}} — font 12, spacing 1.5, align right ──
p4 = doc.paragraphs[4]
for r in p4.runs:
    r.font.size = Pt(12)
    r.bold = False
pf = p4.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ── 4. All table cells — font 12, no bold for value columns ──
for tbl in doc.tables:
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(12)
                    # label column (0) stays bold; value column (1) no bold
                    if ci == 1:
                        r.bold = False

# ── 5. Footer — ensure centered ──
for section in doc.sections:
    footer = section.footer
    for para in footer.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs:
            r.font.size = Pt(10)

doc.save(path)
print('Done OK')
