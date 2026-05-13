from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

src = r'c:/Users/Fonz/ynshouf-app/ynshouf-app/ynshouf-app/YNSHOUF_Template.docx'
out = r'c:/Users/Fonz/ynshouf-app/ynshouf-app/ynshouf-app/YNSHOUF_Template_new.docx'

doc = Document(src)

# ── 1. Fix footer: single clean run with ASCII markers ─────────────────────
for section in doc.sections:
    footer = section.footer
    for para in footer.paragraphs:
        p = para._p
        for child in list(p):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('r', 'hyperlink', 'ins', 'del'):
                p.remove(child)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(
            'YNSHOUF_NAME , מספר עובד YNSHOUF_NUM'
            ', פקח נגב מערבי, ינשוו"פ, רט"ג'
        )
        run.font.name = 'Arial'
        run.font.size = Pt(10)

# ── 2. Add signature section at end of body ────────────────────────────────
h = doc.add_paragraph('חתימה ואישור', style='TiMi Heading 1')
h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

sp = doc.add_paragraph('{{signature}}', style='TiMi Normal')
sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.save(out)
print('Saved to:', out)
print('Close Word and rename manually, or close Word and re-run with save to original path.')
